"""
Генерирует НевзоровПавелЛаб1.docx из report.md,
используя титульную страницу из lab.docx.
"""
import copy
import re
import sys
import io
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LAB_DOCX = os.path.join(BASE_DIR, "lab.docx")
REPORT_MD = os.path.join(BASE_DIR, "report.md")
OUT_DOCX  = os.path.join(BASE_DIR, "НевзоровПавелЛаб1.docx")
RESULTS   = os.path.join(BASE_DIR, "results")

# ───────────────────────────────────────────────────────────────
# Helpers
# ───────────────────────────────────────────────────────────────

def set_run_font(run, size_pt=14, bold=None, italic=None, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def add_paragraph(doc, text="", style="Normal", align=None,
                  size_pt=14, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=size_pt, bold=bold, italic=italic)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.oxml.ns.qn('w:lastRenderedPageBreak'))
    # proper page break
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return p

def add_heading(doc, text, level=1):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    style = style_map.get(level, "Heading 1")
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    # keep font consistent with document
    run.font.name = "Times New Roman"
    if level == 1:
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
    else:
        run.font.size = Pt(14)
        run.bold = True
        run.italic = True
    return p

def add_table_from_md(doc, header_row, rows):
    """Creates a Word table from markdown table rows."""
    col_count = len(header_row)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = h.strip()
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    # data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row):
            cells[ci].text = cell_text.strip()
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    return table

def add_code_block(doc, code_text):
    """Add a code block paragraph with Courier New font."""
    lines = code_text.split('\n')
    for line in lines:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    return

def parse_inline(run_or_para, text):
    """
    Add text to a paragraph, handling **bold** and `code` inline markers.
    Returns the paragraph (if para given) for chaining.
    """
    # Simple tokenizer: split on **...** and `...`
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            content = token[2:-2]
            r = run_or_para.add_run(content)
            r.bold = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
        elif token.startswith('`') and token.endswith('`'):
            content = token[1:-1]
            r = run_or_para.add_run(content)
            r.font.name = "Courier New"
            r.font.size = Pt(12)
        else:
            if token:
                r = run_or_para.add_run(token)
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)

# ───────────────────────────────────────────────────────────────
# Parse markdown lines into structured blocks
# ───────────────────────────────────────────────────────────────

def parse_md_blocks(lines):
    """
    Returns list of (type, content) tuples:
      ('h1'|'h2'|'h3', text)
      ('p', text)
      ('bullet', text)
      ('numbered', text)
      ('table', [header_cells], [[row_cells], ...])
      ('code', code_str)
      ('image', path, alt)
      ('hr', '')
      ('blank', '')
    """
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith('#### '):
            blocks.append(('h3', line[5:].strip()))
            i += 1
        elif line.startswith('### '):
            blocks.append(('h3', line[4:].strip()))
            i += 1
        elif line.startswith('## '):
            blocks.append(('h2', line[3:].strip()))
            i += 1
        elif line.startswith('# '):
            blocks.append(('h1', line[2:].strip()))
            i += 1

        # Horizontal rule
        elif re.match(r'^---+\s*$', line):
            blocks.append(('hr', ''))
            i += 1

        # Fenced code block (may be indented inside list items)
        elif line.lstrip().startswith('```'):
            lang = line.lstrip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].lstrip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            if code_lines:
                blocks.append(('code', '\n'.join(code_lines)))

        # Image
        elif line.startswith('!['):
            m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if m:
                blocks.append(('image', m.group(2), m.group(1)))
            i += 1

        # Table
        elif '|' in line and i + 1 < len(lines) and re.match(r'\s*\|[\s\-:|]+\|\s*$', lines[i + 1]):
            # header row
            header = [c for c in line.split('|') if c.strip() != '']
            i += 1  # skip separator row
            i += 1
            data_rows = []
            while i < len(lines) and '|' in lines[i]:
                row = [c for c in lines[i].split('|') if c != '']
                if row:
                    data_rows.append(row)
                i += 1
            blocks.append(('table', header, data_rows))

        # Bullet list
        elif re.match(r'^[\-\*] ', line):
            blocks.append(('bullet', line[2:].strip()))
            i += 1

        # Numbered list
        elif re.match(r'^\d+\. ', line):
            blocks.append(('numbered', re.sub(r'^\d+\.\s*', '', line).strip()))
            i += 1

        # Blank line
        elif line.strip() == '':
            blocks.append(('blank', ''))
            i += 1

        # Normal paragraph — do NOT merge lines; one para per line
        else:
            blocks.append(('p', line))
            i += 1

    return blocks


# ───────────────────────────────────────────────────────────────
# Main
# ───────────────────────────────────────────────────────────────

import docx

def main():
    # 1. Open lab.docx and clone the entire document to start fresh
    doc = Document(LAB_DOCX)

    # 2. Modify title page text
    paras = doc.paragraphs

    # Fix lab number: paragraph [10] "Лабораторная работа №3" → "№1"
    p10 = paras[10]
    for run in p10.runs:
        if '3' in run.text and 'работа' not in run.text.lower():
            run.text = run.text.replace('3', '1')
            break

    # Fix theme: paragraph [12] "Тема: «Алгоритмы линейной алгебры»"
    p12 = paras[12]
    for run in p12.runs:
        run.text = ''
    if p12.runs:
        p12.runs[0].text = 'Тема: «Определить строку с максимальной суммой элементов»'
    else:
        run = p12.add_run('Тема: «Определить строку с максимальной суммой элементов»')
        set_run_font(run, size_pt=14)

    # Fix author: paragraph [20] "Ф.И.О." → actual name
    p20 = paras[20]
    for run in p20.runs:
        run.text = ''
    if p20.runs:
        p20.runs[0].text = 'Невзоров Павел Евгеньевич'
    else:
        run = p20.add_run('Невзоров Павел Евгеньевич')
        set_run_font(run, size_pt=14)

    # Fix group: paragraph [19] "студент группы "
    p19 = paras[19]
    for run in p19.runs:
        if 'группы' in run.text:
            run.text = 'студент группы ИВТ-22'
            break

    # Fix year: paragraph [27] "Тверь 202 " → "Тверь 2026"
    p27 = paras[27]
    for run in p27.runs:
        if '202' in run.text:
            run.text = run.text.replace('202 ', '2026').replace('202\xa0', '2026')
            break

    # 3. Remove all body content after title page (keep elements 0-27)
    body = doc.element.body
    # Collect elements to remove (index 28 onwards, except last sectPr)
    children = list(body)
    # Find the final sectPr (must be preserved)
    final_sectPr = None
    for child in reversed(children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            final_sectPr = child
            break

    # Remove everything from index 28 onwards (except sectPr)
    for child in children[28:]:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            body.remove(child)

    # 4. Add page break after title
    pb_p = OxmlElement('w:p')
    pb_r = OxmlElement('w:r')
    pb_br = OxmlElement('w:br')
    pb_br.set(qn('w:type'), 'page')
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    # insert before sectPr
    if final_sectPr is not None:
        body.insert(list(body).index(final_sectPr), pb_p)
    else:
        body.append(pb_p)

    # ── helper to add a paragraph in the correct position (before sectPr) ──
    def insert_para(text="", style="Normal", align=None, size_pt=14,
                    bold=False, italic=False, indent_cm=None):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')

        # style
        pStyle = OxmlElement('w:pStyle')
        # map our style names to valid IDs
        style_id_map = {
            "Normal": "Normal",
            "Heading 1": "1",
            "Heading 2": "2",
            "Heading 3": "3",
            "List Paragraph": "ListParagraph",
        }
        pStyle.set(qn('w:val'), style_id_map.get(style, style))
        pPr.append(pStyle)

        # alignment
        if align is not None:
            jc = OxmlElement('w:jc')
            align_map = {
                WD_ALIGN_PARAGRAPH.CENTER: 'center',
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
            }
            jc.set(qn('w:val'), align_map.get(align, 'left'))
            pPr.append(jc)

        # indent
        if indent_cm is not None:
            ind = OxmlElement('w:ind')
            twips = int(indent_cm * 567)  # 1cm = 567 twips
            ind.set(qn('w:left'), str(twips))
            pPr.append(ind)

        if len(pPr):
            p.append(pPr)

        if text:
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            fn = OxmlElement('w:rFonts')
            fn.set(qn('w:ascii'), 'Times New Roman')
            fn.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(fn)

            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(size_pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(size_pt * 2)))
            rPr.append(szCs)

            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)
            if italic:
                i_el = OxmlElement('w:i')
                rPr.append(i_el)

            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
            p.append(r)

        if final_sectPr is not None:
            body.insert(list(body).index(final_sectPr), p)
        else:
            body.append(p)
        return p

    def insert_image(img_path, width_cm=15):
        """Insert image paragraph before sectPr."""
        if not os.path.exists(img_path):
            insert_para(f"[Рисунок: {os.path.basename(img_path)}]",
                        align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12, italic=True)
            return

        # We need to use doc.add_picture but that appends to end.
        # Instead, add a paragraph normally and embed via python-docx run.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        # Move this paragraph before sectPr
        p_elem = p._element
        body.remove(p_elem)
        if final_sectPr is not None:
            body.insert(list(body).index(final_sectPr), p_elem)
        else:
            body.append(p_elem)

    def insert_table_md(header_row, data_rows):
        """Insert a markdown table before sectPr."""
        col_count = len(header_row)
        tbl = doc.add_table(rows=1 + len(data_rows), cols=col_count)
        tbl.style = "Table Grid"

        # header
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(header_row):
            hdr_cells[i].text = h.strip()
            for para in hdr_cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

        # data rows
        for ri, row in enumerate(data_rows):
            cells = tbl.rows[ri + 1].cells
            for ci, cell_text in enumerate(row[:col_count]):
                cells[ci].text = cell_text.strip()
                for para in cells[ci].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)

        # Move table before sectPr
        tbl_elem = tbl._element
        body.remove(tbl_elem)
        if final_sectPr is not None:
            body.insert(list(body).index(final_sectPr), tbl_elem)
        else:
            body.append(tbl_elem)

    def insert_code(code_text):
        lines_code = code_text.split('\n')
        for line in lines_code:
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '567')
            pPr.append(ind)
            p.append(pPr)

            if line:
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                fn = OxmlElement('w:rFonts')
                fn.set(qn('w:ascii'), 'Courier New')
                fn.set(qn('w:hAnsi'), 'Courier New')
                rPr.append(fn)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '20')  # 10pt
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '20')
                rPr.append(szCs)
                r.append(rPr)
                t = OxmlElement('w:t')
                t.text = line
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                r.append(t)
                p.append(r)

            if final_sectPr is not None:
                body.insert(list(body).index(final_sectPr), p)
            else:
                body.append(p)

    def insert_inline_para(text, style_id="Normal", align=None,
                           size_pt=14, base_bold=False, indent_cm=None):
        """Paragraph supporting **bold** and `code` inline markup."""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_id)
        pPr.append(pStyle)

        if align:
            jc = OxmlElement('w:jc')
            align_map = {
                WD_ALIGN_PARAGRAPH.CENTER: 'center',
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
            }
            jc.set(qn('w:val'), align_map.get(align, 'both'))
            pPr.append(jc)

        if indent_cm is not None:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(int(indent_cm * 567)))
            pPr.append(ind)

        p.append(pPr)

        tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)', text)
        for token in tokens:
            if not token:
                continue
            is_bold = base_bold
            is_italic = False
            is_code = False
            content = token
            if token.startswith('**') and token.endswith('**'):
                content = token[2:-2]
                is_bold = True
            elif token.startswith('*') and token.endswith('*') and len(token) > 2:
                content = token[1:-1]
                is_italic = True
            elif token.startswith('`') and token.endswith('`'):
                content = token[1:-1]
                is_code = True

            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            fn = OxmlElement('w:rFonts')
            fname = 'Courier New' if is_code else 'Times New Roman'
            fn.set(qn('w:ascii'), fname)
            fn.set(qn('w:hAnsi'), fname)
            rPr.append(fn)
            fsize = 20 if is_code else int(size_pt * 2)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(fsize))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(fsize))
            rPr.append(szCs)
            if is_bold:
                b = OxmlElement('w:b')
                rPr.append(b)
            if is_italic:
                it = OxmlElement('w:i')
                rPr.append(it)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = content
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
            p.append(r)

        if final_sectPr is not None:
            body.insert(list(body).index(final_sectPr), p)
        else:
            body.append(p)
        return p

    # 5. Parse report.md
    with open(REPORT_MD, encoding='utf-8') as f:
        md_lines = f.read().splitlines()

    # Skip the title block (lines before first "## 1.")
    start = 0
    for idx, line in enumerate(md_lines):
        if line.startswith('## 1.') or line.startswith('## Оглавление'):
            start = idx
            break

    # Skip the ToC section entirely (## Оглавление ... until next ## N.)
    # We'll regenerate no ToC — just start from "## 1."
    for idx in range(start, len(md_lines)):
        if re.match(r'^## 1\.', md_lines[idx]):
            start = idx
            break

    blocks = parse_md_blocks(md_lines[start:])

    # 6. Render blocks into the document
    in_appendix = False
    prev_was_blank = False

    for block in blocks:
        btype = block[0]

        if btype == 'blank':
            prev_was_blank = True
            continue

        if btype == 'hr':
            prev_was_blank = False
            continue

        if btype == 'h1':
            # h1 (#) — skip, used only as document title (already in title page)
            prev_was_blank = False

        elif btype == 'h2':
            # h2 (##) — main sections → Heading 1
            text = block[1]
            text = re.sub(r'^\d+\.\s+', '', text)  # strip "1. "
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), '1')  # Heading 1
            pPr.append(pStyle)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            fn = OxmlElement('w:rFonts')
            fn.set(qn('w:ascii'), 'Times New Roman')
            fn.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(fn)
            b = OxmlElement('w:b')
            rPr.append(b)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '32')  # 16pt
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            r.append(t)
            p.append(r)
            if final_sectPr is not None:
                body.insert(list(body).index(final_sectPr), p)
            else:
                body.append(p)
            if 'Приложени' in text:
                in_appendix = True
            prev_was_blank = False

        elif btype == 'h3':
            # h3 (###) — subsections → Heading 2
            text = block[1]
            text = re.sub(r'^\d+\.\d+\.?\s+', '', text)  # strip "4.1 "
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), '2')  # Heading 2
            pPr.append(pStyle)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            fn = OxmlElement('w:rFonts')
            fn.set(qn('w:ascii'), 'Times New Roman')
            fn.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(fn)
            b = OxmlElement('w:b')
            rPr.append(b)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '28')  # 14pt
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            r.append(t)
            p.append(r)
            if final_sectPr is not None:
                body.insert(list(body).index(final_sectPr), p)
            else:
                body.append(p)
            prev_was_blank = False

        elif btype == 'p':
            text = block[1]
            # Skip markdown-only lines
            if text.startswith('!['):
                continue
            # Skip raw LaTeX formulas
            if text.strip().startswith('$$') or text.strip().startswith('$'):
                continue
            # Skip pure markdown link lines like "[text](#anchor)"
            if re.match(r'^\[.+\]\(#.+\)\s*$', text.strip()):
                continue
            insert_inline_para(text, style_id="Normal",
                               align=WD_ALIGN_PARAGRAPH.JUSTIFY, size_pt=14)
            prev_was_blank = False

        elif btype == 'bullet':
            insert_inline_para(block[1], style_id="ListParagraph",
                               align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               size_pt=14, indent_cm=1.0)
            prev_was_blank = False

        elif btype == 'numbered':
            insert_inline_para(block[1], style_id="ListParagraph",
                               align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               size_pt=14, indent_cm=1.0)
            prev_was_blank = False

        elif btype == 'table':
            header_row = block[1]
            data_rows  = block[2]
            insert_table_md(header_row, data_rows)
            insert_para('')  # spacing after table
            prev_was_blank = False

        elif btype == 'code':
            code_text = block[1]
            if in_appendix:
                # Appendix code: smaller font, with border shading effect via indent
                insert_code(code_text)
            else:
                insert_code(code_text)
            prev_was_blank = False

        elif btype == 'image':
            img_rel_path = block[1]
            # Resolve path relative to BASE_DIR
            img_path = os.path.join(BASE_DIR, img_rel_path.lstrip('./'))
            insert_image(img_path, width_cm=14)
            prev_was_blank = False

    # 7. Save
    doc.save(OUT_DOCX)
    print(f"Saved: {OUT_DOCX}")


if __name__ == '__main__':
    main()
